#!/usr/bin/env python3
"""Generate the COMP8347 Week 4 Guestbook report as a polished .docx.

Run from anywhere:  uv run python gen_report_docx.py
Output: report.docx (screenshots pulled from ./screenshots/).
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, "screenshots")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x2E, 0x5E, 0x8C)
LIGHT = "E8F0F8"
HEADER_FILL = "1F3A5F"
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x22, 0x22, 0x22)


def shade(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    sd = OxmlElement("w:shd")
    sd.set(qn("w:val"), "clear")
    sd.set(qn("w:fill"), color_hex)
    tcPr.append(sd)


def set_cell(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E5E8C")
    pbdr.append(bottom)
    pPr.append(pbdr)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BLUE


def body(doc, text, size=10.5, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = DARK


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead)
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(rest)
            r2.font.size = Pt(10.5)
        else:
            r = p.add_run(it)
            r.font.size = Pt(10.5)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, HEADER_FILL)
        set_cell(c, htext, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            set_cell(cells[ci], str(val), size=9.5)
            if ri % 2 == 1:
                shade(cells[ci], LIGHT)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)


def figure(doc, filename, caption, width=6.2):
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        body(doc, f"[missing screenshot: {filename}]", size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = GREY


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# Cover
for _ in range(4):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("COMP8347"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Internet Applications / Distributed Systems"); r.font.size = Pt(13); r.font.color.rgb = BLUE
doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Week 4 Class Activity (28.05)"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DARK
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Challenge: Fixing the Guestbook App"); r.font.size = Pt(13); r.font.color.rgb = GREY
for _ in range(8):
    doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Individual Assignment — Report"); r.font.size = Pt(11); r.font.color.rgb = GREY
doc.add_page_break()

# 1. Problem Overview
h1(doc, "1. Problem Overview")
h2(doc, "a) What was wrong in the provided Guestbook app")
body(doc, "The starter project was a Django app (guestbook) that was intentionally broken in nine places. "
          "None of the four required features worked end to end. The defects were:")
table(doc,
      ["#", "Location", "Defect"],
      [
          ["1", "forms.py", "Meta.fields only had 'message', so display_name was never rendered or saved."],
          ["2", "forms.py", "clean_message() returned the raw value, allowing empty / whitespace-only messages."],
          ["3", "views.py (add_entry)", "Session counter was re-assigned to its current value instead of incremented."],
          ["4", "views.py (add_entry)", "Cookie pre-fill set initial for a display_name field the form did not have."],
          ["5", "views.py (set_name_cookie)", "Cookie written as 'display_name' but read as 'displayname' (key mismatch)."],
          ["6", "views.py (clear_name_cookie)", "Deleted a non-existent key 'display_namex', leaving the real cookie."],
          ["7", "entry_list.html", "'Set my name cookie' link sent no name parameter, so the value was blank."],
          ["8", "add_entry.html", "POST form missing CSRF token, so submissions failed with HTTP 403."],
          ["9", "add_entry.html", "Template rendered form.display_name, a field the form did not define."],
      ],
      widths=[0.4, 1.9, 3.9])

h2(doc, "b) The required user features (in my own words)")
bullets(doc, [
    ("Add an entry — ", "a visitor types a display name and a message and submits them."),
    ("List entries — ", "all messages are shown, each with its author and the time it was posted."),
    ("Remember the name — ", "a cookie stores the visitor's name so the form is pre-filled next time."),
    ("Count this session's entries — ", "the page shows how many entries the visitor added this session."),
])

# 2. Implementation Summary
h1(doc, "2. Implementation Summary (After Fix)")

h2(doc, "A. Add name + message")
body(doc, "UI description. The 'Sign the Guestbook' page (/guestbook/add/) has a single-line Display name input "
          "and a multi-line Message textarea, plus a Save button. Each field shows a placeholder stating its maximum "
          "length (40 and 280 characters). The HTML maxlength attribute was deliberately removed so the user can type "
          "past the limit and the server-side validator can reject the input with a clear message.")
body(doc, "What happens on submit. The form POSTs to add_entry. The view builds an EntryForm from request.POST; if it "
          "is valid it saves the Entry, increments the session counter, sets the name cookie, and redirects to the list "
          "(Post/Redirect/Get). If invalid, it re-renders the page with inline error messages.")
body(doc, "Validation and error handling:")
bullets(doc, [
    ("Empty name/message — ", "clean_message() strips whitespace and rejects empty input; the textarea is also "
                                    "required, so an empty message is never saved."),
    ("Overly long inputs — ", "limits are 40 characters for the name and 280 for the message (matching the model). "
                                    "Longer input is rejected server-side."),
])
figure(doc, "A1-add-form.png", "A1 — Add-entry form filled in (Display name + Message).")
figure(doc, "A2-add-success.png", "A2 — Entry added successfully; session counter shows 1.")
figure(doc, "A3-error-empty.png", "A3 — Empty message is blocked on submit.")
figure(doc, "A4-error-toolong.png", "A4 — Over-long message input.")

h2(doc, "B. List messages with time + author")
body(doc, "Data model (guestbook/models.py, class Entry):")
table(doc,
      ["Field", "Type", "Notes"],
      [
          ["display_name", "CharField(max_length=40, blank=True, default=\"\")", "author; 'Anonymous' if blank"],
          ["message", "TextField(max_length=280)", "the message body"],
          ["created_at", "DateTimeField(auto_now_add=True)", "timestamp set on creation"],
      ],
      widths=[1.4, 2.9, 1.9])
bullets(doc, [
    ("Storage — ", "entries are kept in the starter app's SQLite database (db.sqlite3) via the Django ORM."),
    ("Sorting order — ", "newest first (order_by('-created_at')); a guestbook should surface the latest activity at the top."),
    ("Time formatting — ", "Django filter date:\"M d, Y H:i\", e.g. May 30, 2026 01:46. TIME_ZONE='UTC' and "
                                 "USE_TZ=True, so times are in UTC."),
])
figure(doc, "B1-list-multiple.png", "B1 — Multiple messages listed with author and time, newest first.")

h2(doc, "C. Cookie to remember the user's name")
bullets(doc, [
    ("Cookie name / value — ", "'displayname', storing the visitor's display name as a string."),
    ("When set/updated — ", "on every successful submission (using the posted name), and via the 'Set my name cookie' "
                                  "form on the list page."),
    ("Expiry policy — ", "persistent cookie, max_age = 7 days. Persistent (not session) because the goal is to "
                               "remember the visitor across visits."),
    ("How the UI uses it — ", "on GET, the view passes the cookie as the form's initial value so the Display name "
                                    "field is auto-filled; the list page also shows the saved name."),
])
figure(doc, "C1-cookie-set.png", "C1 — DevTools showing the 'displayname' cookie set after the first submit.")
figure(doc, "C2-name-prefilled.png", "C2 — Add form auto-filled with the remembered name.")

h2(doc, "D. Track how many entries added this session")
bullets(doc, [
    ("Definition of 'session' — ", "a server-side session (Django session framework, keyed by the sessionid cookie)."),
    ("Where the counter is stored — ", "the server session dictionary: request.session['gb_added']."),
    ("When it increments — ", "only on a successful POST, after is_valid() and save(): "
                                    "session['gb_added'] = get('gb_added', 0) + 1."),
    ("Refresh — ", "preserved (same session cookie)."),
    ("New tab (same browser) — ", "continues (same session cookie sent)."),
    ("New browser / incognito — ", "a fresh session, so the counter restarts at 0."),
])
figure(doc, "D1-counter-1.png", "D1 — Session entry counter.")
figure(doc, "D2-counter-2.png", "D2 — Counter increases as more entries are added.")

h1(doc, "3. Verification")
body(doc, "An automated test suite (guestbook/tests.py, 5 tests) confirms the behaviour: adding an entry redirects, "
          "increments the counter and sets the name cookie; the list shows author + message, newest first; the add form "
          "pre-fills the name from the cookie; and empty or over-long messages are rejected and not saved. All 5 tests "
          "pass with: python manage.py test guestbook.")

out = os.path.join(HERE, "report.docx")
doc.save(out)
print("Saved:", out)
